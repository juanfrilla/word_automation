from docx import Document
from datetime import datetime
import locale
import re
import textwrap


def datetime_to_dateformat(input_date: datetime, format: str):
    locale.setlocale(locale.LC_TIME, "es_ES.UTF-8")
    return input_date.strftime(format)


def remove_blank_pages(doc: Document) -> Document:
    sections_to_remove = []
    for i, section in enumerate(doc.sections):
        if (
            not section.footer.paragraphs
            and not section.header.paragraphs
            and not section.paragraphs
        ):
            sections_to_remove.append(i)
    for i in reversed(sections_to_remove):
        doc.element.body.remove(doc.sections[i]._element)
    return doc

def find_and_complete_dots(doc:Document) -> Document:
    dot_pattern = re.compile(r'\.{4,}')
    raw_paragraphs = doc.paragraphs
    raw_paragraphs_text = [p.text for p in raw_paragraphs]
    

    for paragraph in raw_paragraphs:
        print(paragraph.text)
        
            
    return doc

# def fill_line_with_dots(doc: Document) -> Document:
#     for paragraph in doc.paragraphs:
#         dots_at_end = 0
#         for char in reversed(paragraph.text):
#             if char == ".":
#                 dots_at_end += 1
#             else:
#                 break

#         # Add dots to fill the line until the end
#         if dots_at_end >= 4:
#             paragraph.text += "." * (80 - len(paragraph.text))  # Assuming a line length of 80 characters
#     return doc

# def last_line_of_paragraph_has_four_dots(paragraph):
    
#     dot_pattern = r'\.{4,}'
#     match = re.search(dot_pattern, line)
#     if match:
#         return True
#     return False

# def fill_lines_with_dots(doc: Document):

#     for paragraph in doc.paragraphs:
#         #max_line_length = max(len(line) for line in paragraph)
#         if last_line_of_paragraph_has_four_dots(paragraph):
#                 filled_lines = ['.'.join([line, '.' * (max_line_length - len(line))]) for line in lines]
#                 paragraph.text = '\n'.join(filled_lines)
#         else:
#             continue
#     return doc
